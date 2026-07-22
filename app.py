"""
Servidor web principal para la Suite OSINT de Extracción de Usuarios.
"""
__author__ = "Jorge R."
__copyright__ = "Copyright 2026, Proyecto DataExtractor"
__license__ = "MIT"
__version__ = "2.0.0"

from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
import threading
import uuid
import os
import json
import datetime
import bcrypt
from extractors import extractor
from extractors import extractor_selenium
from extractors import extractor_tiktok
from extractors import extractor_x
from extractors import extractor_fb
from extractors import extractor_linkedin
from extractors import extractor_youtube
from extractors import osint_search

app = Flask(__name__)

# Configurar Flask-Limiter para proteger contra fuerza bruta
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"]
)

tasks = {}

HISTORY_FILE = 'history.json'
SETTINGS_FILE = 'settings.json'

def load_json(filepath, default):
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return default
    return default

def save_json(filepath, data):
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def add_to_history(task_id, target, method, count, out_path, platform="Instagram"):
    history = load_json(HISTORY_FILE, [])
    history.append({
        "id": task_id,
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "target": target,
        "method": method,
        "platform": platform,
        "count": count,
        "file": out_path
    })
    save_json(HISTORY_FILE, history)
    # Auto Telegram notification
    try:
        settings = load_json(SETTINGS_FILE, {})
        if settings.get('telegram_auto') and settings.get('telegram_bot_token') and settings.get('telegram_chat_id'):
            import requests
            msg = (f"📊 *DataExtractor - Nueva Extracción*\n\n"
                   f"👤 Objetivo: `{target}`\n"
                   f"📱 Plataforma: {platform}\n"
                   f"📋 Método: {method}\n"
                   f"📈 Registros: {count}\n"
                   f"🕐 Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            requests.post(
                f"https://api.telegram.org/bot{settings['telegram_bot_token']}/sendMessage",
                json={"chat_id": settings['telegram_chat_id'], "text": msg, "parse_mode": "Markdown"},
                timeout=5
            )
    except Exception:
        pass  # Don't break extraction if Telegram fails

@app.route('/api/history', methods=['GET'])
def get_history():
    return jsonify(load_json(HISTORY_FILE, []))

@app.route('/api/history/<item_id>', methods=['DELETE'])
def delete_history_item(item_id):
    history = load_json(HISTORY_FILE, [])
    history = [h for h in history if h.get('id') != item_id]
    save_json(HISTORY_FILE, history)
    return jsonify({"success": True})

@app.route('/api/history', methods=['DELETE'])
def clear_history():
    save_json(HISTORY_FILE, [])
    return jsonify({"success": True})

@app.route('/api/settings', methods=['GET', 'POST'])
def manage_settings():
    if request.method == 'POST':
        save_json(SETTINGS_FILE, request.json)
        return jsonify({"status": "success"})
    return jsonify(load_json(SETTINGS_FILE, {
        "defaultUser": "", "defaultPassword": "",
        "ig_user": "", "ig_pass": "",
    }))

@app.route('/api/export', methods=['GET'])
def export_data():
    """Export all data (history, settings, tags) as a single JSON backup."""
    backup = {
        "history": load_json(HISTORY_FILE, []),
        "settings": load_json(SETTINGS_FILE, {}),
        "tags": load_json('tags.json', {}),
        "export_date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "version": "2.0"
    }
    return jsonify(backup)

@app.route('/api/import', methods=['POST'])
def import_data():
    """Import data from a backup JSON file."""
    data = request.json
    if not data:
        return jsonify({"error": "No data provided"}), 400
    try:
        if 'history' in data:
            save_json(HISTORY_FILE, data['history'])
        if 'settings' in data:
            save_json(SETTINGS_FILE, data['settings'])
        if 'tags' in data:
            save_json('tags.json', data['tags'])
        return jsonify({"success": True, "message": "Datos importados correctamente"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

TAGS_FILE = 'tags.json'

@app.route('/api/tags', methods=['GET'])
def get_tags():
    return jsonify(load_json(TAGS_FILE, {}))

@app.route('/api/tags', methods=['POST'])
def add_tag():
    data = request.json
    username = data.get('username', '')
    label = data.get('label', '')
    color = data.get('color', 'blue')
    if not username or not label:
        return jsonify({"error": "username and label required"}), 400
    tags = load_json(TAGS_FILE, {})
    if username not in tags:
        tags[username] = []
    # Avoid duplicates
    if not any(t['label'] == label for t in tags[username]):
        tags[username].append({"label": label, "color": color})
    save_json(TAGS_FILE, tags)
    return jsonify({"success": True})

@app.route('/api/tags/<username>/<label>', methods=['DELETE'])
def delete_tag(username, label):
    tags = load_json(TAGS_FILE, {})
    if username in tags:
        tags[username] = [t for t in tags[username] if t['label'] != label]
        if not tags[username]:
            del tags[username]
        save_json(TAGS_FILE, tags)
    return jsonify({"success": True})

@app.route('/')
def index():
    return render_template('index.html')

NOTES_FILE = 'notes.json'

@app.route('/api/dashboard-stats')
def dashboard_stats():
    """Compute dashboard statistics from history."""
    history = load_json(HISTORY_FILE, [])
    total = len(history)
    users = set()
    platforms = set()
    total_records = 0
    monthly = {}

    for item in history:
        users.add(item.get('target', ''))
        platforms.add(item.get('platform', 'osint'))
        total_records += item.get('count', 0)
        date_str = item.get('date', '')
        if date_str:
            month_key = date_str[:7]  # "2026-07"
            monthly[month_key] = monthly.get(month_key, 0) + 1

    # Last 6 months
    sorted_months = sorted(monthly.keys())[-6:]
    monthly_activity = [{'month': m, 'count': monthly[m]} for m in sorted_months]

    # Recent 5
    recent = history[-5:][::-1]

    return jsonify({
        'total': total,
        'unique_users': len(users),
        'platforms': len(platforms),
        'total_records': total_records,
        'monthly_activity': monthly_activity,
        'recent': recent
    })

@app.route('/api/notes/<username>', methods=['GET'])
def get_note(username):
    notes = load_json(NOTES_FILE, {})
    return jsonify({'username': username, 'note': notes.get(username, '')})

@app.route('/api/notes', methods=['POST'])
def save_note():
    data = request.json
    username = data.get('username', '')
    note = data.get('note', '')
    notes = load_json(NOTES_FILE, {})
    if note:
        notes[username] = note
    elif username in notes:
        del notes[username]
    save_json(NOTES_FILE, notes)
    return jsonify({'success': True})

# =================================================================
# ENDPOINT: Búsqueda OSINT (Rastrear perfil en múltiples redes)
# =================================================================
@app.route('/osint/search', methods=['POST'])
def osint_search_endpoint():
    data = request.json
    username = data.get('username', '').strip()
    redes = data.get('platforms', ['instagram', 'tiktok', 'x', 'facebook'])

    if not username:
        return jsonify({"error": "Username is required"}), 400

    task_id = str(uuid.uuid4())
    out_path = os.path.join("Output", f"osint_{task_id[:8]}.csv")

    tasks[task_id] = {
        'status': 'running',
        'logs': [],
        'progress': 0,
        'out_path': out_path,
        'driver': None,
        'osint_results': None,
        'type': 'osint'
    }

    def log_cb(m):
        tasks[task_id]['logs'].append(m)

    def prog_cb(c, u):
        tasks[task_id]['progress'] = c

    def run_osint():
        try:
            results = osint_search.buscar_en_todas_las_redes(
                username, redes=redes, log_callback=log_cb, progress_callback=prog_cb
            )
            tasks[task_id]['osint_results'] = results
            osint_search.guardar_resultados_osint(results, out_path, log_cb)
            tasks[task_id]['status'] = 'completed'
            add_to_history(task_id, username, 'OSINT Multi-Red',
                           results.get('total_found', 0), out_path, 'Multi-Platform')
        except Exception as e:
            log_cb(f"Error OSINT: {e}")
            tasks[task_id]['status'] = 'error'

    threading.Thread(target=run_osint, daemon=True).start()
    return jsonify({"task_id": task_id})

@app.route('/osint/results/<task_id>')
def osint_results(task_id):
    if task_id in tasks and tasks[task_id].get('osint_results'):
        return jsonify(tasks[task_id]['osint_results'])
    return jsonify({"error": "No results"}), 404

# =================================================================
# ENDPOINT: Extracción Principal (Basado en Selenium, todas las redes)
# =================================================================
LOGIN_URLS = {
    'instagram': 'https://www.instagram.com/accounts/login/',
    'tiktok': 'https://www.tiktok.com/login',
    'x': 'https://x.com/i/flow/login',
    'facebook': 'https://www.facebook.com/login',
    'linkedin': 'https://www.linkedin.com/login',
    'youtube': 'https://accounts.google.com/ServiceLogin?service=youtube',
}

@app.route('/start', methods=['POST'])
def start():
    data = request.json
    method = data.get('method')
    target = data.get('target', '').strip()
    platform = data.get('platform', 'instagram')
    extract_type = data.get('extract_type', 'followers')
    depth = data.get('depth', 'basic')

    if method == 'Selenium_Scan':
        task_id = data.get('task_id')
        if task_id not in tasks:
            return jsonify({"error": "Task not found"}), 404

        driver = tasks[task_id]['driver']
        out_path = tasks[task_id]['out_path']

        def run_selenium():
            tasks[task_id]['status'] = 'running'
            def log_cb(m): tasks[task_id]['logs'].append(m)
            def prog_cb(c, u): tasks[task_id]['progress'] = c

            log_cb(f"=== EXTRACCIÓN: {platform.upper()} ===")
            log_cb(f"Objetivo: @{target}")
            log_cb(f"Tipo: {extract_type} | Profundidad: {depth}")
            log_cb(f"{'='*40}")

            try:
                if platform == 'instagram':
                    from extractors.extractor_selenium import extraer_lista_ig
                    extraer_lista_ig(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
                elif platform == 'x':
                    from extractors.extractor_x import extraer_lista_x
                    extraer_lista_x(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
                elif platform == 'facebook':
                    from extractors.extractor_fb import extraer_lista_fb
                    extraer_lista_fb(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
                elif platform == 'tiktok':
                    from extractors.extractor_tiktok import extraer_lista_tiktok
                    extraer_lista_tiktok(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
                elif platform == 'linkedin':
                    from extractors.extractor_linkedin import extraer_lista_linkedin
                    extraer_lista_linkedin(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
                elif platform == 'youtube':
                    from extractors.extractor_youtube import extraer_lista_youtube
                    extraer_lista_youtube(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
            except Exception as e:
                log_cb(f"Error durante la extracción: {e}")

            log_cb("--- FIN DEL PROCESO ---")
            tasks[task_id]['status'] = 'completed'
            add_to_history(task_id, target, f'Selenium ({extract_type}/{depth})',
                           tasks[task_id]['progress'], out_path, platform.upper())

        threading.Thread(target=run_selenium, daemon=True).start()
        return jsonify({"task_id": task_id})

    # ---- Selenium_Prepare: Open browser at login page ----
    if method == 'Selenium_Prepare':
        task_id = str(uuid.uuid4())
        out_path = os.path.join("Output", f"extract_{task_id[:8]}.csv")
        
        # Ensure Output directory exists
        os.makedirs("Output", exist_ok=True)

        tasks[task_id] = {
            'status': 'starting',
            'logs': [],
            'progress': 0,
            'out_path': out_path,
            'driver': None,
            'type': 'extraction',
            'extract_type': extract_type,
            'depth': depth,
        }

        def log_callback(msg):
            tasks[task_id]['logs'].append(msg)

        def run_selenium_prepare():
            try:
                driver = extractor_selenium.configurar_driver(log_callback)
                tasks[task_id]['driver'] = driver

                # Navigate to login page
                login_url = LOGIN_URLS.get(platform, 'about:blank')
                log_callback(f"Navegando a la página de login de {platform.upper()}...")
                driver.get(login_url)

                tasks[task_id]['status'] = 'browser_ready'
                log_callback(f"✓ Navegador abierto en {platform.upper()}.")
                log_callback(f"→ Inicie sesión con su cuenta.")
                log_callback(f"→ Una vez logueado, pulse 'Comenzar Extracción Automática'.")
            except Exception as e:
                log_callback(f"Error al abrir navegador: {e}")
                tasks[task_id]['status'] = 'error'

        threading.Thread(target=run_selenium_prepare, daemon=True).start()
        return jsonify({"task_id": task_id})

    return jsonify({"error": "Método no soportado"}), 400

@app.route('/status/<task_id>')
def status(task_id):
    if task_id in tasks:
        task = tasks[task_id]
        logs = list(task['logs'])
        task['logs'].clear()
        return jsonify({
            'status': task['status'],
            'logs': logs,
            'progress': task['progress']
        })
    return jsonify({'status': 'not_found'}), 404

@app.route('/download/<task_id>')
def download(task_id):
    fmt = request.args.get('format', 'csv')

    # Try in-memory tasks first, then fall back to history.json
    csv_path = None
    if task_id in tasks and os.path.exists(tasks[task_id].get('out_path', '')):
        csv_path = tasks[task_id]['out_path']
    else:
        history = load_json(HISTORY_FILE, [])
        for item in history:
            if item.get('id') == task_id and os.path.exists(item.get('file', '')):
                csv_path = item['file']
                break

    if not csv_path:
        return jsonify({"error": "Archivo no encontrado"}), 404

    if fmt == 'json':
        import csv as csv_mod
        rows = []
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv_mod.DictReader(f)
            for row in reader:
                rows.append(row)
        json_path = csv_path.replace('.csv', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump({"total": len(rows), "data": rows,
                        "exported_at": datetime.datetime.now().isoformat()}, f, indent=2, ensure_ascii=False)
        return send_file(json_path, as_attachment=True, download_name=f"data_{task_id[:6]}.json")

    elif fmt == 'excel':
        import csv as csv_mod
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "Datos Extraídos"

        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv_mod.reader(f)
            headers = next(reader)

            # Style header
            header_font = Font(bold=True, color="FFFFFF", size=11)
            header_fill = PatternFill(start_color="1a1a2e", end_color="1a1a2e", fill_type="solid")
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header.upper())
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border

            for row_idx, row in enumerate(reader, 2):
                for col, val in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col, value=val)
                    cell.border = thin_border

            # Auto-width columns
            for col in ws.columns:
                max_len = max(len(str(c.value or "")) for c in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

        xlsx_path = csv_path.replace('.csv', '.xlsx')
        wb.save(xlsx_path)
        return send_file(xlsx_path, as_attachment=True, download_name=f"data_{task_id[:6]}.xlsx")

    # Default: CSV
    return send_file(csv_path, as_attachment=True, download_name=f"data_{task_id[:6]}.csv")


# =================================================================
# ENDPOINT: Cross-platform analysis
# =================================================================
@app.route('/api/cross-analysis', methods=['POST'])
def cross_analysis():
    """Compara usernames entre múltiples archivos CSV para encontrar coincidencias."""
    import csv as csv_mod

    if 'files' not in request.files:
        return jsonify({"error": "No files uploaded"}), 400

    files = request.files.getlist('files')
    if len(files) < 2:
        return jsonify({"error": "Se necesitan al menos 2 archivos"}), 400

    platform_users = {}
    for f in files:
        name = f.filename or "unknown"
        reader = csv_mod.DictReader(f.stream.read().decode('utf-8').splitlines())
        users = set()
        for row in reader:
            uname = row.get('username', row.get('Username', '')).strip().lower()
            if uname:
                users.add(uname)
        platform_users[name] = users

    # Find common users across all platforms
    all_sets = list(platform_users.values())
    common_all = set.intersection(*all_sets) if all_sets else set()

    # Pairwise intersections
    pairwise = {}
    keys = list(platform_users.keys())
    for i in range(len(keys)):
        for j in range(i + 1, len(keys)):
            common = platform_users[keys[i]] & platform_users[keys[j]]
            pairwise[f"{keys[i]} ∩ {keys[j]}"] = list(common)

    return jsonify({
        "platforms": {k: len(v) for k, v in platform_users.items()},
        "common_all": list(common_all),
        "common_all_count": len(common_all),
        "pairwise": {k: {"count": len(v), "users": v[:200]} for k, v in pairwise.items()},
    })


# =================================================================
# ENDPOINT: Descargar Fotos de Perfil
# =================================================================
@app.route('/api/download-photos/<task_id>', methods=['POST'])
def download_photos(task_id):
    """Descarga las fotos de perfil de los usuarios extraídos."""
    import csv as csv_mod
    import zipfile
    import io

    if task_id not in tasks or not os.path.exists(tasks[task_id]['out_path']):
        return jsonify({"error": "Task not found"}), 404

    csv_path = tasks[task_id]['out_path']
    photo_urls = []

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv_mod.DictReader(f)
        for row in reader:
            pic = row.get('profile_pic', '')
            uname = row.get('username', '')
            if pic and pic.startswith('http'):
                photo_urls.append((uname, pic))

    if not photo_urls:
        return jsonify({"error": "No profile photos found in the data"}), 404

    # Download and zip photos
    import requests as req
    zip_buffer = io.BytesIO()
    downloaded = 0

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for uname, url in photo_urls[:500]:  # Limit to 500
            try:
                resp = req.get(url, timeout=10)
                if resp.status_code == 200:
                    ext = '.jpg'
                    if 'png' in resp.headers.get('content-type', ''):
                        ext = '.png'
                    zf.writestr(f"{uname}{ext}", resp.content)
                    downloaded += 1
            except Exception:
                continue

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True,
                     download_name=f"fotos_{task_id[:6]}.zip",
                     mimetype='application/zip')


# =================================================================
# ENDPOINT: Generar Informe en PDF
# =================================================================
@app.route('/api/report/<task_id>')
def generate_report(task_id):
    """Genera un informe PDF profesional OSINT."""
    import csv as csv_mod
    from fpdf import FPDF

    # Try in-memory tasks first, then fall back to history.json
    csv_path = None
    if task_id in tasks and os.path.exists(tasks[task_id].get('out_path', '')):
        csv_path = tasks[task_id]['out_path']
    else:
        # Search in persistent history
        history = load_json(HISTORY_FILE, [])
        for item in history:
            if item.get('id') == task_id and os.path.exists(item.get('file', '')):
                csv_path = item['file']
                break

    if not csv_path:
        return jsonify({"error": "No se encontró el archivo CSV para esta tarea. Es posible que los datos hayan sido eliminados."}), 404

    def safe_text(text, max_len=25):
        """Sanitize text for PDF (remove non-latin1 characters)."""
        s = str(text)[:max_len]
        # Replace characters outside latin-1 range
        return ''.join(c if ord(c) < 256 else '?' for c in s)

    rows = []
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv_mod.DictReader(f)
            headers = reader.fieldnames or []
            for row in reader:
                rows.append(row)
    except Exception as e:
        return f"Error reading CSV: {e}", 500

    # Check format requested
    fmt = request.args.get('format', 'pdf')

    if fmt == 'docx':
        # === Word Export ===
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles['Normal']
            style.font.size = Pt(10)

            # Title
            title = doc.add_heading('Informe OSINT - DataExtractor', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Generado: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").font.size = Pt(9)
            p.add_run(f"Total registros: {len(rows)}").font.size = Pt(9)

            doc.add_paragraph()

            # Table
            if headers and rows:
                display_headers = headers[:8]
                table = doc.add_table(rows=1, cols=len(display_headers))
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                for i, h in enumerate(display_headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(h)[:30]
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(8)

                # Data rows (limit 500)
                for row in rows[:500]:
                    row_cells = table.add_row().cells
                    for i, h in enumerate(display_headers):
                        row_cells[i].text = str(row.get(h, ''))[:40]
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(7)

                if len(rows) > 500:
                    doc.add_paragraph(f'... y {len(rows) - 500} registros más')

            docx_path = csv_path.replace('.csv', '_informe.docx')
            doc.save(docx_path)
            return send_file(docx_path, as_attachment=True, download_name=f"informe_{task_id[:6]}.docx")
        except Exception as e:
            return f"Error generating Word: {e}", 500

    # === PDF Export (default) ===
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(0, 15, "Informe OSINT - DataExtractor", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, f"Generado: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                 new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.cell(0, 8, f"Task ID: {task_id[:8]}... | Total registros: {len(rows)}",
                 new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(10)

        # Table header
        if headers and rows:
            display_headers = headers[:6]
            col_w = 190 / len(display_headers)

            pdf.set_font("Helvetica", "B", 8)
            pdf.set_fill_color(26, 26, 46)
            pdf.set_text_color(255, 255, 255)
            for h in display_headers:
                pdf.cell(col_w, 8, safe_text(h, 20), border=1, fill=True, align="C")
            pdf.ln()

            # Table rows
            pdf.set_font("Helvetica", "", 7)
            pdf.set_text_color(0, 0, 0)
            for row in rows[:500]:
                for h in display_headers:
                    val = safe_text(row.get(h, ""), 25)
                    pdf.cell(col_w, 6, val, border=1)
                pdf.ln()

            if len(rows) > 500:
                pdf.ln(5)
                pdf.set_font("Helvetica", "I", 9)
                pdf.cell(0, 8, f"... y {len(rows) - 500} registros mas (ver CSV/Excel para datos completos)")

        pdf_path = csv_path.replace('.csv', '_informe.pdf')
        pdf.output(pdf_path)
        return send_file(pdf_path, as_attachment=True, download_name=f"informe_{task_id[:6]}.pdf")
    except Exception as e:
        return f"Error generating PDF: {e}", 500


# =================================================================
# ENDPOINT: Scheduled Extractions (APScheduler)
# =================================================================
from apscheduler.schedulers.background import BackgroundScheduler

scheduler = BackgroundScheduler()
scheduler.start()
SCHEDULED_FILE = 'scheduled_tasks.json'


def load_scheduled():
    if os.path.exists(SCHEDULED_FILE):
        with open(SCHEDULED_FILE, 'r') as f:
            return json.load(f)
    return []


def save_scheduled(data):
    with open(SCHEDULED_FILE, 'w') as f:
        json.dump(data, f, indent=2)


def run_scheduled_extraction(job_id, platform, target, extract_type, depth):
    """Runs a scheduled extraction using headless Selenium."""
    import extractor_selenium

    task_id = f"sched_{job_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
    out_path = os.path.join('output', f'{task_id}.csv')
    os.makedirs('output', exist_ok=True)

    tasks[task_id] = {
        'status': 'running',
        'logs': [],
        'progress': 0,
        'out_path': out_path,
        'type': 'scheduled',
    }

    def log_cb(m):
        tasks[task_id]['logs'].append(m)
        print(f"[SCHED {job_id}] {m}")

    def prog_cb(c, u):
        tasks[task_id]['progress'] = c

    log_cb(f"=== EXTRACCIÓN PROGRAMADA: {platform.upper()} ===")
    log_cb(f"Objetivo: @{target} | Tipo: {extract_type} | Profundidad: {depth}")

    try:
        driver = extractor_selenium.configurar_driver(log_cb)
        login_url = LOGIN_URLS.get(platform, 'about:blank')
        driver.get(login_url)

        # In scheduled mode, we try headless; the user should have saved cookies
        import time
        time.sleep(5)

        if platform == 'instagram':
            from extractor_selenium import extraer_lista_ig
            extraer_lista_ig(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
        elif platform == 'x':
            from extractor_x import extraer_lista_x
            extraer_lista_x(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
        elif platform == 'facebook':
            from extractor_fb import extraer_lista_fb
            extraer_lista_fb(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
        elif platform == 'tiktok':
            from extractor_tiktok import extraer_lista_tiktok
            extraer_lista_tiktok(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
        elif platform == 'linkedin':
            from extractor_linkedin import extraer_lista_linkedin
            extraer_lista_linkedin(driver, target, out_path, extract_type, depth, log_cb, prog_cb)
        elif platform == 'youtube':
            from extractor_youtube import extraer_lista_youtube
            extraer_lista_youtube(driver, target, out_path, extract_type, depth, log_cb, prog_cb)

        driver.quit()
    except Exception as e:
        log_cb(f"Error: {e}")

    tasks[task_id]['status'] = 'completed'
    add_to_history(task_id, target, f'Programada ({extract_type}/{depth})',
                   tasks[task_id]['progress'], out_path, platform.upper())


@app.route('/api/scheduled', methods=['GET'])
def get_scheduled():
    """Returns all scheduled jobs."""
    jobs_data = load_scheduled()
    # Enrich with next run time
    for job_data in jobs_data:
        job = scheduler.get_job(job_data.get('job_id'))
        if job:
            job_data['next_run'] = str(job.next_run_time) if job.next_run_time else 'N/A'
            job_data['active'] = True
        else:
            job_data['active'] = False
    return jsonify(jobs_data)


@app.route('/api/scheduled', methods=['POST'])
def create_scheduled():
    """Create a new scheduled extraction."""
    data = request.json
    platform = data.get('platform', 'instagram')
    target = data.get('target', '')
    extract_type = data.get('extract_type', 'followers')
    depth = data.get('depth', 'basic')
    interval_days = int(data.get('interval_days', 7))

    job_id = str(uuid.uuid4())[:8]

    scheduler.add_job(
        run_scheduled_extraction,
        'interval',
        days=interval_days,
        id=job_id,
        args=[job_id, platform, target, extract_type, depth],
        name=f"{platform}/{target}",
    )

    job_record = {
        'job_id': job_id,
        'platform': platform,
        'target': target,
        'extract_type': extract_type,
        'depth': depth,
        'interval_days': interval_days,
        'created_at': datetime.datetime.now().isoformat(),
    }

    jobs = load_scheduled()
    jobs.append(job_record)
    save_scheduled(jobs)

    return jsonify({"success": True, "job_id": job_id})


@app.route('/api/scheduled/<job_id>', methods=['DELETE'])
def delete_scheduled(job_id):
    """Delete a scheduled job."""
    try:
        scheduler.remove_job(job_id)
    except Exception:
        pass

    jobs = load_scheduled()
    jobs = [j for j in jobs if j.get('job_id') != job_id]
    save_scheduled(jobs)
    return jsonify({"success": True})


# =================================================================
# ENDPOINT: Reverse Search by Email
# =================================================================
@app.route('/api/reverse-search', methods=['POST'])
def reverse_search():
    """Búsqueda inversa: dado un email o teléfono, busca perfiles asociados."""
    import requests as req

    data = request.json
    query = data.get('query', '').strip()
    if not query:
        return jsonify({"error": "Se necesita un email o teléfono"}), 400

    results = []
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }

    # Check if it looks like an email
    is_email = '@' in query

    platforms_checked = []

    # Instagram: Try password reset flow (publicly accessible)
    try:
        platforms_checked.append('Instagram')
        r = req.post('https://www.instagram.com/accounts/account_recovery_send_ajax/',
                      data={'email_or_username': query},
                      headers={**headers, 'X-Requested-With': 'XMLHttpRequest',
                               'Referer': 'https://www.instagram.com/'},
                      timeout=10)
        if r.status_code == 200:
            body = r.text.lower()
            if 'sent' in body or 'email' in body or 'enviado' in body:
                results.append({
                    'platform': 'Instagram',
                    'found': True,
                    'method': 'password_reset',
                    'detail': 'Cuenta asociada encontrada (se envió email de recuperación).'
                })
    except Exception:
        pass

    # Facebook: Graph search (public)
    if is_email:
        try:
            platforms_checked.append('Facebook')
            r = req.get(f'https://www.facebook.com/login/identify/?ctx=recover&email={query}',
                        headers=headers, timeout=10, allow_redirects=False)
            if r.status_code in (200, 302):
                if 'identify' not in r.headers.get('location', '').lower():
                    results.append({
                        'platform': 'Facebook',
                        'found': True,
                        'method': 'recovery_flow',
                        'detail': 'Cuenta asociada detectada.'
                    })
        except Exception:
            pass

    # X/Twitter: Check signup
    try:
        platforms_checked.append('X (Twitter)')
        r = req.get(f'https://api.twitter.com/i/users/email_available.json?email={query}',
                     headers=headers, timeout=10)
        if r.status_code == 200:
            data_resp = r.json()
            if not data_resp.get('valid', True):
                results.append({
                    'platform': 'X (Twitter)',
                    'found': True,
                    'method': 'email_check',
                    'detail': 'Email ya registrado en X.'
                })
    except Exception:
        pass

    return jsonify({
        'query': query,
        'type': 'email' if is_email else 'phone',
        'platforms_checked': platforms_checked,
        'results': results,
        'total_found': len(results)
    })


# =================================================================
# ENDPOINT: Búsqueda Sherlock (Rastrear en más de 120 plataformas web)
# =================================================================
from extractors import sherlock_checker

@app.route('/api/sherlock', methods=['POST'])
def sherlock_search():
    data = request.json
    username = data.get('username', '').strip().lstrip('@')
    if not username:
        return jsonify({"error": "Se necesita un username"}), 400

    task_id = str(uuid.uuid4())
    tasks[task_id] = {
        'status': 'running',
        'logs': [],
        'progress': 0,
        'total': sherlock_checker.get_platform_count(),
        'type': 'sherlock',
        'results': None,
    }

    def run_sherlock():
        def log_cb(m):
            tasks[task_id]['logs'].append(m)
        def prog_cb(checked, total, platform, found):
            tasks[task_id]['progress'] = checked

        from extractors import sherlock_checker
        result = sherlock_checker.check_username(username, log_callback=log_cb, progress_callback=prog_cb)
        tasks[task_id]['results'] = result
        tasks[task_id]['status'] = 'completed'

    threading.Thread(target=run_sherlock, daemon=True).start()
    return jsonify({"task_id": task_id, "total": tasks[task_id]['total']})


@app.route('/api/sherlock/<task_id>', methods=['GET'])
def sherlock_status(task_id):
    t = tasks.get(task_id)
    if not t:
        return jsonify({"error": "Not found"}), 404
    return jsonify({
        "status": t['status'],
        "progress": t['progress'],
        "total": t.get('total', 0),
        "results": t.get('results'),
    })


# =================================================================
# ENDPOINT: Bulk Search (multiple usernames)
# =================================================================
@app.route('/api/bulk-search', methods=['POST'])
def bulk_search():
    data = request.json
    usernames = data.get('usernames', [])
    redes = data.get('redes', None)

    if not usernames:
        return jsonify({"error": "No usernames provided"}), 400

    task_id = str(uuid.uuid4())
    tasks[task_id] = {
        'status': 'running',
        'logs': [],
        'progress': 0,
        'total': len(usernames),
        'type': 'bulk',
        'results': {},
    }

    def run_bulk():
        for idx, username in enumerate(usernames, 1):
            username = username.strip().lstrip('@')
            if not username:
                continue
            tasks[task_id]['logs'].append(f"\n=== [{idx}/{len(usernames)}] Buscando @{username} ===")

            def log_cb(m):
                tasks[task_id]['logs'].append(m)
            def prog_cb(c, p):
                pass

            result = osint_search.buscar_en_todas_las_redes(
                username, redes=redes, log_callback=log_cb, progress_callback=prog_cb
            )
            tasks[task_id]['results'][username] = result
            tasks[task_id]['progress'] = idx

        tasks[task_id]['status'] = 'completed'

    threading.Thread(target=run_bulk, daemon=True).start()
    return jsonify({"task_id": task_id, "total": len(usernames)})


@app.route('/api/bulk-search/<task_id>', methods=['GET'])
def bulk_status(task_id):
    t = tasks.get(task_id)
    if not t:
        return jsonify({"error": "Not found"}), 404
    return jsonify({
        "status": t['status'],
        "progress": t['progress'],
        "total": t.get('total', 0),
        "results": t.get('results'),
        "logs": t.get('logs', [])[-50:],
    })


# =================================================================
# ENDPOINT: Detección de Bots y Perfiles Falsos
# =================================================================

@app.route('/api/bot-analysis/<task_id>', methods=['GET'])
def bot_analysis(task_id):
    t = tasks.get(task_id)
    if not t:
        return jsonify({"error": "Task not found"}), 404
    csv_path = t.get('out_path', '')
    if not csv_path or not os.path.exists(csv_path):
        return jsonify({"error": "CSV not found"}), 404
    from extractors import bot_detector
    result = bot_detector.analyze_csv(csv_path)
    return jsonify(result)


# =================================================================
# ENDPOINT: Detección de Cambios (Comparar extracciones anteriores)
# =================================================================
@app.route('/api/diff', methods=['POST'])
def diff_extractions():
    import csv as csv_mod
    data = request.json
    task_id_old = data.get('task_id_old', '')
    task_id_new = data.get('task_id_new', '')

    # Get CSV paths from tasks
    t_old = tasks.get(task_id_old)
    t_new = tasks.get(task_id_new)

    csv_old = t_old.get('out_path', '') if t_old else ''
    csv_new = t_new.get('out_path', '') if t_new else ''

    # Also try history
    if not csv_old or not os.path.exists(csv_old):
        history = load_json(HISTORY_FILE, [])
        for h in history:
            if h.get('task_id') == task_id_old:
                csv_old = h.get('file_path', '')
                break
    if not csv_new or not os.path.exists(csv_new):
        history = load_json(HISTORY_FILE, [])
        for h in history:
            if h.get('task_id') == task_id_new:
                csv_new = h.get('file_path', '')
                break

    if not csv_old or not csv_new or not os.path.exists(csv_old) or not os.path.exists(csv_new):
        return jsonify({"error": "No se encontraron los archivos CSV"}), 404

    def read_users(path):
        users = {}
        with open(path, mode='r', encoding='utf-8') as f:
            reader = csv_mod.DictReader(f)
            for row in reader:
                uname = row.get('username', '').strip()
                if uname:
                    users[uname] = row
        return users

    old_users = read_users(csv_old)
    new_users = read_users(csv_new)

    old_set = set(old_users.keys())
    new_set = set(new_users.keys())

    gained = list(new_set - old_set)
    lost = list(old_set - new_set)
    maintained = list(old_set & new_set)

    # Detect changes in maintained users
    changes = []
    for uname in maintained:
        old_data = old_users[uname]
        new_data = new_users[uname]
        user_changes = []
        for key in new_data:
            if key != 'username' and old_data.get(key, '') != new_data.get(key, ''):
                user_changes.append({
                    "field": key,
                    "old": old_data.get(key, ''),
                    "new": new_data.get(key, ''),
                })
        if user_changes:
            changes.append({"username": uname, "changes": user_changes})

    return jsonify({
        "old_total": len(old_users),
        "new_total": len(new_users),
        "gained": gained,
        "lost": lost,
        "maintained": len(maintained),
        "changes": changes,
        "gained_count": len(gained),
        "lost_count": len(lost),
        "changes_count": len(changes),
    })


# =================================================================
# ENDPOINT: Análisis de Geolocalización
# =================================================================
GEOCACHE_FILE = 'geocache.json'

def load_geocache():
    if os.path.exists(GEOCACHE_FILE):
        with open(GEOCACHE_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_geocache(cache):
    with open(GEOCACHE_FILE, 'w', encoding='utf-8') as f:
        json.dump(cache, f, indent=2)

@app.route('/api/geodata/<task_id>', methods=['GET'])
def geodata(task_id):
    import csv as csv_mod
    from geopy.geocoders import Nominatim
    import time as time_mod

    t = tasks.get(task_id)
    csv_path = ''
    if t:
        csv_path = t.get('out_path', '')

    # Also try history
    if not csv_path or not os.path.exists(csv_path):
        history = load_json(HISTORY_FILE, [])
        for h in history:
            if h.get('task_id') == task_id:
                csv_path = h.get('file_path', '')
                break

    if not csv_path or not os.path.exists(csv_path):
        return jsonify({"error": "CSV not found"}), 404

    cache = load_geocache()
    geolocator = Nominatim(user_agent="dataextractor_osint", timeout=5)
    markers = []

    with open(csv_path, mode='r', encoding='utf-8') as f:
        reader = csv_mod.DictReader(f)
        for row in reader:
            location = row.get('location', '').strip()
            if not location or len(location) < 2:
                continue

            username = row.get('username', '')

            if location in cache:
                coords = cache[location]
                if coords:
                    markers.append({
                        "username": username,
                        "location": location,
                        "lat": coords[0],
                        "lng": coords[1],
                    })
            else:
                try:
                    result = geolocator.geocode(location)
                    if result:
                        cache[location] = [result.latitude, result.longitude]
                        markers.append({
                            "username": username,
                            "location": location,
                            "lat": result.latitude,
                            "lng": result.longitude,
                        })
                    else:
                        cache[location] = None
                    time_mod.sleep(1)  # Nominatim rate limit
                except Exception:
                    cache[location] = None

    save_geocache(cache)
    return jsonify({"markers": markers, "total": len(markers)})






# =================================================================
# ENDPOINT: Notificaciones de Telegram
# =================================================================
@app.route('/api/telegram/test', methods=['POST'])
def test_telegram():
    import requests as req
    data = request.json
    bot_token = data.get('bot_token', '')
    chat_id = data.get('chat_id', '')

    if not bot_token or not chat_id:
        return jsonify({"error": "Bot token y chat ID son necesarios"}), 400

    try:
        url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
        resp = req.post(url, json={
            "chat_id": chat_id,
            "text": "DataExtractor OSINT: Conexion exitosa! Las notificaciones estan configuradas.",
            "parse_mode": "HTML"
        }, timeout=10)
        if resp.status_code == 200:
            # Save to settings
            settings = load_json(SETTINGS_FILE, {})
            settings['telegram_bot_token'] = bot_token
            settings['telegram_chat_id'] = chat_id
            save_json(SETTINGS_FILE, settings)
            return jsonify({"success": True, "message": "Mensaje enviado correctamente"})
        else:
            return jsonify({"error": f"Error de Telegram: {resp.text}"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def send_telegram_notification(message):
    """Send a notification via Telegram if configured."""
    import requests as req
    settings = load_json(SETTINGS_FILE, {})
    bot_token = settings.get('telegram_bot_token', '')
    chat_id = settings.get('telegram_chat_id', '')
    if not bot_token or not chat_id:
        return
    try:
        url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
        req.post(url, json={
            "chat_id": chat_id,
            "text": message,
            "parse_mode": "HTML"
        }, timeout=10)
    except Exception:
        pass


# =================================================================
# ENDPOINT: Autenticación (Flask-Login)
# =================================================================
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login_page'

# Generar o cargar una clave secreta local persistente para las sesiones
secret_key_file = '.flask_secret'
if not os.path.exists(secret_key_file):
    with open(secret_key_file, 'wb') as f:
        f.write(os.urandom(24))
with open(secret_key_file, 'rb') as f:
    local_secret = f.read()

app.secret_key = os.environ.get('SECRET_KEY', local_secret)

# Seguridad: Proteger cookies contra ataques CSRF e inyección JS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_HTTPONLY'] = True

# Session expires when browser closes by default.
# Only persists if "remember me" is checked (30 days).
app.config['SESSION_PERMANENT'] = False
app.config['REMEMBER_COOKIE_DURATION'] = datetime.timedelta(days=30)

USERS_FILE = 'users.json'

class User(UserMixin):
    def __init__(self, username):
        self.id = username

@login_manager.user_loader
def load_user(username):
    users = load_json(USERS_FILE, {})
    if username in users:
        return User(username)
    return None

@app.route('/login', methods=['GET'])
def login_page():
    users = load_json(USERS_FILE, {})
    if not users:
        # No users yet, skip auth
        return render_template('index.html')
    return render_template('login.html')

@app.route('/api/auth/login', methods=['POST'])
@limiter.limit("5 per minute")
def api_login():
    data = request.json
    username = data.get('username', '')
    password = data.get('password', '')
    remember = data.get('remember', False)

    users = load_json(USERS_FILE, {})
    if username in users:
        stored_hash = users[username].get('password', '')
        if bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8')):
            login_user(User(username), remember=remember)
            return jsonify({"success": True})

    return jsonify({"error": "Credenciales incorrectas"}), 401

@app.route('/api/auth/register', methods=['POST'])
def api_register():
    data = request.json
    username = data.get('username', '')
    password = data.get('password', '')

    if not username or not password or len(password) < 4:
        return jsonify({"error": "Username y password (min 4 chars) necesarios"}), 400

    users = load_json(USERS_FILE, {})
    if username in users:
        return jsonify({"error": "El usuario ya existe"}), 400

    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    users[username] = {"password": hashed, "created": datetime.datetime.now().isoformat()}
    save_json(USERS_FILE, users)
    login_user(User(username))
    return jsonify({"success": True})

@app.route('/api/auth/logout', methods=['POST'])
def api_logout():
    logout_user()
    return jsonify({"success": True})

@app.route('/api/auth/status', methods=['GET'])
def auth_status():
    users = load_json(USERS_FILE, {})
    return jsonify({
        "authenticated": current_user.is_authenticated if hasattr(current_user, 'is_authenticated') else False,
        "username": current_user.id if hasattr(current_user, 'id') and current_user.is_authenticated else None,
        "auth_enabled": len(users) > 0,
    })


# =================================================================
# ENDPOINT: Analíticas e Intersección de Datos (Encontrar usuarios en común)
# =================================================================
@app.route('/api/v1/docs', methods=['GET'])
def api_docs():
    docs = {
        "name": "DataExtractor OSINT API",
        "version": "2.0",
        "endpoints": [
            {"method": "POST", "path": "/osint/search", "desc": "Buscar username en 6 redes sociales"},
            {"method": "GET", "path": "/osint/results/<task_id>", "desc": "Obtener resultados OSINT"},
            {"method": "POST", "path": "/start", "desc": "Iniciar extraccion Selenium"},
            {"method": "GET", "path": "/status/<task_id>", "desc": "Estado de tarea"},
            {"method": "GET", "path": "/download/<task_id>?format=csv|excel|json", "desc": "Descargar resultados"},
            {"method": "GET", "path": "/api/report/<task_id>", "desc": "Generar informe PDF"},
            {"method": "POST", "path": "/api/cross-analysis", "desc": "Analisis cruzado multi-plataforma"},
            {"method": "POST", "path": "/api/sherlock", "desc": "Sherlock: comprobar username en 120+ plataformas"},
            {"method": "GET", "path": "/api/sherlock/<task_id>", "desc": "Estado/resultados Sherlock"},
            {"method": "POST", "path": "/api/bulk-search", "desc": "Busqueda en lote (multiples usernames)"},
            {"method": "GET", "path": "/api/bulk-search/<task_id>", "desc": "Estado/resultados bulk search"},
            {"method": "GET", "path": "/api/bot-analysis/<task_id>", "desc": "Analisis de bots en extraccion"},
            {"method": "POST", "path": "/api/diff", "desc": "Comparar dos extracciones (diff)"},
            {"method": "GET", "path": "/api/geodata/<task_id>", "desc": "Datos de geolocalizacion"},
            {"method": "GET/POST", "path": "/api/tags", "desc": "Sistema de etiquetas"},
            {"method": "GET/POST/DELETE", "path": "/api/scheduled", "desc": "Tareas programadas"},
            {"method": "POST", "path": "/api/reverse-search", "desc": "Busqueda inversa email/telefono"},
            {"method": "POST", "path": "/api/telegram/test", "desc": "Probar notificacion Telegram"},
            {"method": "POST", "path": "/api/auth/login", "desc": "Login"},
            {"method": "POST", "path": "/api/auth/register", "desc": "Registrar usuario"},
            {"method": "GET", "path": "/api/auth/status", "desc": "Estado de autenticacion"},
            {"method": "GET", "path": "/api/v1/docs", "desc": "Esta documentacion"},
        ]
    }
    return jsonify(docs)

# =================================================================
# ENDPOINT: Media Extraction (photos, videos, likes)
# =================================================================
@app.route('/api/media-extract', methods=['POST'])
def start_media_extraction():
    """Start a background media extraction task."""
    data = request.json
    platform = data.get('platform', 'instagram')
    username = data.get('username', '').strip().lstrip('@')
    include_likes = data.get('include_likes', False)
    max_posts = min(int(data.get('max_posts', 30)), 100)

    if not username:
        return jsonify({"error": "Username required"}), 400

    task_id = f"media_{uuid.uuid4().hex[:8]}"
    tasks[task_id] = {
        'status': 'running',
        'logs': [],
        'progress': 0,
        'total': max_posts,
        'results': None,
        'type': 'media'
    }

    def log_cb(m):
        tasks[task_id]['logs'].append(m)

    def prog_cb(current, total):
        tasks[task_id]['progress'] = current
        tasks[task_id]['total'] = total

    def run_extraction():
        import media_extractor
        try:
            if platform == 'instagram':
                settings = load_json(SETTINGS_FILE, {})
                results = media_extractor.extract_instagram_media(
                    username, 
                    session_user=settings.get('ig_user'),
                    session_file=f'session_{settings.get("ig_user", "")}' if settings.get('ig_user') else None,
                    include_likes=include_likes,
                    max_posts=max_posts,
                    log_callback=log_cb,
                    progress_callback=prog_cb
                )
            elif platform == 'tiktok':
                results = media_extractor.extract_tiktok_media(
                    username, max_videos=max_posts,
                    log_callback=log_cb, progress_callback=prog_cb
                )
            elif platform == 'youtube':
                results = media_extractor.extract_youtube_media(
                    username, max_videos=max_posts,
                    log_callback=log_cb, progress_callback=prog_cb
                )
            else:
                results = {'posts': [], 'errors': [f'Plataforma {platform} no soportada para media']}

            tasks[task_id]['results'] = results
            tasks[task_id]['status'] = 'completed'
            log_cb(f"Extraccion completada: {len(results.get('posts', []))} elementos")
        except Exception as e:
            tasks[task_id]['status'] = 'error'
            log_cb(f"Error: {str(e)}")

    thread = threading.Thread(target=run_extraction, daemon=True)
    thread.start()

    return jsonify({"task_id": task_id})


@app.route('/api/media-extract/<task_id>', methods=['GET'])
def get_media_status(task_id):
    if task_id not in tasks:
        return jsonify({"error": "Task not found"}), 404
    t = tasks[task_id]
    return jsonify({
        "status": t['status'],
        "progress": t.get('progress', 0),
        "total": t.get('total', 0),
        "logs": t.get('logs', [])[-20:],
        "results": t.get('results')
    })


# Serve media files
@app.route('/media/<path:filename>')
def serve_media(filename):
    return send_from_directory('media', filename)

if __name__ == '__main__':
    # Por defecto debug está desactivado por seguridad
    is_debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(host='0.0.0.0', debug=is_debug, port=5000)
